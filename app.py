import streamlit as st
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- SAHIFANI SOZLASH ---
st.set_page_config(page_title="Cloud Baza 2026", layout="wide", page_icon="📝")

# --- GOOGLE SHEETS ULANISH ---
try:
    conn = st.connection("gsheets", type=GSheetsConnection)
except:
    st.error("Secrets sozlamalarini tekshiring!")
    st.stop()

# --- LOGIN TIZIMI ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    st.title("🔐 Tizimga kirish")
    user = st.text_input("Login:")
    pas = st.text_input("Parol:", type="password")
    if st.button("Kirish"):
        if user == "admin" and pas == "12345":
            st.session_state['logged_in'] = True
            st.rerun()
        else:
            st.error("Xato!")
    st.stop()

# --- TO'LIQ SHARTNOMA GENERATORI ---
def generate_official_contract(d):
    doc = Document()
    
    # Standart shrift: Times New Roman, 11pt
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    def add_centered_bold(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True

    def add_justified(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)

    # 1-SAHIFA: SARLAVHA
    add_centered_bold("Махсулот қийматини бўлиб тўлаш шарти билан тузилган\n№ " + d['nomer'] + "- сонли олди сотди\nШАРТНОМА")
    doc.add_paragraph(d['sana']).alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run("Мен ")
    intro.add_run(d['ism']).bold = True
    intro.add_run(f" Узбекистон Фукароси, паспорт № {d['pasport']} {d['pas_sana']} йилда {d['pas_joy']} томонидан берилган, {d['manzil']} истиқомат қилувчи, телефон {d['tel']} «Харидор» бир тарафдан ва OOO \"NEW DREAMS STAR\" номидан Устав асосида фаолият юритувчи ва кейинги ўринларда “Сотувчи” деб номланувчи директор Нурбеков У.Ю. иккинчи тарафдан ушбу шартномани қуйидагилар ҳақида туздик:")

    # BANDLAR 1-4
    sections_1_4 = [
        ("1. Шартнома предмети", 
         "1.1. Ушбу Шартномага асосан “Сотувчи” товарларни “Харидор”нинг эгалигига топшириш, “Харидор” эса ушбу товарларни қабул қилиб олиш ва улар учун белгиланган қийматни бўлиб тўлаш мажбуриятини ўз зиммаларига оладилар.\n1.2. Товарлар харидорга тўлик топширилган вақтдан бошлаб, унинг қиймати тўлиқ тўланишига қадар, сотилган товарлар харидорнинг қарзини тўлаш мажбуриятини бажаришини таъминlash учун сотувчи томонидан гаровга олинган деб тан олинади."),
        ("2. Шартнома суммаси ва ҳисоб-китоблар тартиби", 
         "2.1. Ушбу Шартноманинг суммаси 1-иловада.\n2.2. Товар “Харидор”га муддатли тўлов шартларида топширилади.\n2.4. Харидор товарни қабул қилишдан асоссиз бош тортса, аванс тўловининг 50% миқдорида жарима тўлайди."),
        ("3. Товарни тақдим қилиш тартиби", "3.1. Сотувчи ҳужжатлар расмийлаштирилгач товарни етказиб беради."),
        ("4. Товарларга тўлов киритиш тартиби", f"4.1. Товарларга тўлов Харидор томонидан 2-иловадаги жадвал асосида амалга оширилади. 4.4. Тўланган пуллар аввало жарима тўловига, сўнгра қарздорликни қоплашга йўналтирилади.")
    ]
    for title, text in sections_1_4:
        add_centered_bold(title)
        add_justified(text)

    # 2-3 SAHIFALAR (BANDLAR 5-9)
    doc.add_page_break()
    sections_5_9 = [
        ("5. Сотувчининг назорати", "5.1. Харидор паспорт маълумотлари, яшаш манзили ёки иш жойи ўзгаргани ҳақида Сотувчини хабардор қилиши шарт."),
        ("6. Харидорнинг мажбуриятларини бажариши кафолатлари", "6.1. Харидорнинг мажбуриятлари бажариши кафолати сифатида кафиллик ёки банк картасидаги маблағлар хизмат қилиши мумкин."),
        ("7. Қарзнинг қолган қисмини муддатдан олдин қайтарилиши", "7.1. Тўлов графиги бузилса ёки Харидорнинг молиявий ҳолати ёмонлашса, Сотувчи қарзни тўлиқ қайтаришни талаб қилишга ҳақли. 7.2. Харидор талабномани олгач 3 кун ичида тўловни амалга ошириши лозим."),
        ("8. Тарафларнинг мажбуриятлари", "8.1. Сотувчи сифатли товар етказиши, 8.2. Харидор эса товарни кўриб қабул қилиши ва вақтида тўлаши шарт."),
        ("9. Тарафларнинг ҳуқуқлари", f"9.1. Сотувчи Харидордан тўлов қобилиятини тасдиқловчи ҳужжатларни талаб қилиш ҳуқуқига эга. 9.1.8. Тўлов кечикса, Сотувчи Харидорнинг маълумотларини Гаров реестрига ёки маҳалла қўмиталарига тақдим қилиши мумкин.")
    ]
    for title, text in sections_5_9:
        add_centered_bold(title)
        add_justified(text)

    # 4-5 SAHIFALAR (BANDLAR 10-15)
    doc.add_page_break()
    sections_10_15 = [
        ("10. Тарафларнинг масъулияти", f"10.5. Тўлов муддати ўтса, Харидор кечиктирилган ҳар бир кун учун 2.0 % жарима тўлайди. 10.8. Сотувчи уяли алоқа воситасини масофадан туриб Идентификатор (Apple ID/Gmail) орқали қулфлаб қўйиш ҳуқуқига эга."),
        ("11. Товарларни топшириш шартлари", "11.2. Товар топширилаётганда унга Сотувчи томонидан Идентификатор ўрнатилади."),
        ("12. Форс-мажор", "12.1. Енгиб бўлмас кучлар таъсирида масъулият чекланади."),
        ("13. Шартномани ўзгартириш ва бекор қилиш", "13.1. Шартномага ўзгартиришлар фақат ёзма равишда киритилади."),
        ("14. Низоларни хал қилиш", "14.1. Низолар Сирдарё туманлараро судларида кўриб чиқилади."),
        ("15. Якуний қоидалар", "15.7. Шартнома 2 нусхада тузилди ва иккаласи ҳам тенг юридик кучга эга.")
    ]
    for title, text in sections_10_15:
        add_centered_bold(title)
        add_justified(text)

    # ILOVALAR SAHIFASI
    doc.add_page_break()
    add_centered_bold("1-илова\nТовар спецификацияси")
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Махсулот", "Микдор", "Нарх", "Сумма"
    row = table1.rows[1].cells
    row[0].text, row[1].text, row[2].text, row[3].text = d['mahsulot'], "1", d['summa'], d['summa']

    doc.add_paragraph(f"\nЖАМИ: {d['summa']} ({d['summa_soz']}) сўм.")

    # GRAFIK SAHIFASI
    doc.add_page_break()
    add_centered_bold("2-илова\nТўловлар жадвали")
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text = "Тўлов тури", "Муддати", "Сумма"
    
    for i in range(1, int(d['oylar']) + 1):
        r = table2.add_row().cells
        r[0].text = f"{i}-тўлов"
        r[1].text = f"27.{i:02d}.2026 гача"
        r[2].text = d['oylik']

    # IMZOLAR
    doc.add_page_break()
    add_centered_bold("ТАРАФЛАРНИНГ ИМЗОЛАРИ")
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.rows[0].cells[0].text = "ХАРИДОР"
    sig_table.rows[0].cells[1].text = "СОТУВЧИ"
    b_row = sig_table.rows[1].cells
    b_row[0].text = f"{d['ism']}\nПаспорт: {d['pasport']}\nТел: {d['tel']}\nМанзил: {d['manzil']}\n\n________ (имзо)"
    b_row[1].text = "OOO 'NEW DREAMS STAR'\nИНН: 306547414\nДиректор: Нурбеков У.Ю.\n\n________ (имзо)"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- STREAMLIT UI ---
st.sidebar.title("🚀 Contract Generator")
tanlov = st.sidebar.radio("Bo'lim:", ["📄 Shartnoma yaratish", "📊 Statistika"])

if tanlov == "📄 Shartnoma yaratish":
    st.header("📝 Rasmiy shartnomani to'ldirish")
    
    with st.form("contract_form"):
        col1, col2 = st.columns(2)
        with col1:
            nomer = st.text_input("Shartnoma №:", "3080")
            sana = st.text_input("Sana:", "27.12.2025")
            ism = st.text_input("F.I.SH:", "URINBAYEV SHOHJAHON SHAROF O’G’LI")
            pas = st.text_input("Pasport №:", "AD6259891")
            pas_sana = st.text_input("Pasport berilgan sana:", "23.02.2024")
        with col2:
            pas_joy = st.text_input("Bergan tashkilot:", "JIZZAX VILOYATI IIV")
            manzil = st.text_area("Mijoz manzili:", "JIZZAX VILOYATI TOSHLOQ QFY")
            tel = st.text_input("Telefon:", "90 487 97 77")
            mahsulot = st.text_input("Mahsulot nomi:", "IPHONE 13 PRO")
            summa = st.text_input("Jami summa:", "5 436 000")
            summa_soz = st.text_input("Summa so'z bilan:", "BESH MILLION TO’RT YUZ O’TIZ OLTI MING")
            oylar = st.selectbox("Muddat (oy):", [3, 6, 9, 12, 24])
            oylik = st.text_input("Oylik to'lov:", "906 000")
        
        submitted = st.form_submit_button("Ma'lumotlarni saqlash va tasdiqlash")

    if submitted:
        data = {
            'nomer': nomer, 'sana': sana, 'ism': ism, 'pasport': pas,
            'pas_sana': pas_sana, 'pas_joy': pas_joy, 'manzil': manzil,
            'tel': tel, 'mahsulot': mahsulot, 'summa': summa,
            'summa_soz': summa_soz, 'oylar': oylar, 'oylik': oylik
        }
        f = generate_official_contract(data)
        st.success("✅ Shartnoma tayyorlandi!")
        st.download_button(
            label="📥 TO'LIQ WORD SHARTNOMANI YUKLAB OLISH",
            data=f,
            file_name=f"Contract_{nomer}_{ism}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
